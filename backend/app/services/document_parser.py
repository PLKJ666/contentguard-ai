"""
文档解析服务
从 PDF/Word/Excel 文档中提取纯文本
"""

import asyncio
import logging
import os
import tempfile
from typing import Optional

import httpx

logger = logging.getLogger(__name__)


class DocumentParser:
    """从文档中提取纯文本"""

    @staticmethod
    async def download_and_parse(document_url: str, document_name: str) -> str:
        """
        下载文档并解析为纯文本

        优先使用 TOS SDK 直接下载（私有桶无需签名），
        回退到 HTTP 预签名 URL 下载。

        Args:
            document_url: 文档 URL (TOS)
            document_name: 原始文件名（用于判断格式）

        Returns:
            提取的纯文本
        """
        tmp_path: Optional[str] = None
        try:
            ext = (
                document_name.rsplit(".", 1)[-1].lower() if "." in document_name else ""
            )

            # 优先用 TOS SDK 直接下载（后端有 AK/SK，无需签名 URL）
            content = await DocumentParser._download_via_tos_sdk(document_url)

            if content is None:
                # 回退：生成预签名 URL 后用 HTTP 下载
                content = await DocumentParser._download_via_signed_url(document_url)

            # 跳过过大的文件（>50MB），解析可能非常慢且阻塞
            if len(content) > 50 * 1024 * 1024:
                logger.warning(
                    f"文件 {document_name} 过大 ({len(content) // 1024 // 1024}MB)，已跳过"
                )
                return ""

            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            # 文件解析可能很慢（CPU 密集），放到线程池执行
            return await asyncio.to_thread(
                DocumentParser.parse_file, tmp_path, document_name
            )
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # 图片提取限制
    MAX_IMAGES = 10
    MAX_IMAGE_SIZE = 2 * 1024 * 1024  # 2MB per image base64

    @staticmethod
    async def download_and_get_images(
        document_url: str, document_name: str
    ) -> Optional[list[str]]:
        """
        下载文档并提取嵌入的图片，返回 base64 编码列表。

        支持格式：
        - PDF: 图片型 PDF 转页面图片
        - DOCX: 提取 word/media/ 中的嵌入图片
        - XLSX: 提取 worksheet 中的嵌入图片

        Returns:
            base64 图片列表，无图片时返回 None
        """
        ext = document_name.rsplit(".", 1)[-1].lower() if "." in document_name else ""
        if ext not in ("pdf", "doc", "docx", "xls", "xlsx"):
            return None

        tmp_path: Optional[str] = None
        try:
            file_content = await DocumentParser._download_via_tos_sdk(document_url)
            if file_content is None:
                file_content = await DocumentParser._download_via_signed_url(
                    document_url
                )

            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            if ext == "pdf":
                if DocumentParser.is_image_pdf(tmp_path):
                    return DocumentParser.pdf_to_images_base64(tmp_path)
                return None
            elif ext in ("doc", "docx"):
                images = await asyncio.to_thread(
                    DocumentParser._extract_docx_images, tmp_path
                )
                return images if images else None
            elif ext in ("xls", "xlsx"):
                images = await asyncio.to_thread(
                    DocumentParser._extract_xlsx_images, tmp_path
                )
                return images if images else None
            return None
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    @staticmethod
    async def _download_via_tos_sdk(document_url: str) -> Optional[bytes]:
        """通过本地存储或 TOS SDK 直接下载文件，在线程池中执行避免阻塞"""

        def _sync_download() -> Optional[bytes]:
            try:
                from app.services.oss import download_from_tos

                data = download_from_tos(document_url)
                if data is not None:
                    logger.info(f"文件直读成功: source={document_url}, size={len(data)}")
                    return data
                logger.debug(f"文件直读失败，准备回退 HTTP: {document_url}")
                return None
            except Exception as e:
                logger.warning(f"文件直读失败，将回退 HTTP: {e}")
                return None

        return await asyncio.to_thread(_sync_download)

    @staticmethod
    async def _download_via_signed_url(document_url: str) -> bytes:
        """生成预签名 URL 后通过 HTTP 下载"""
        from app.services.oss import generate_presigned_url, parse_file_key_from_url

        file_key = parse_file_key_from_url(document_url)
        signed_url = generate_presigned_url(file_key, expire_seconds=300)
        logger.info(f"HTTP 下载: key={file_key}")

        async with httpx.AsyncClient(timeout=60.0) as client:
            resp = await client.get(signed_url)
            resp.raise_for_status()
        logger.info(f"HTTP 下载成功: {len(resp.content)} bytes")
        return resp.content

    @staticmethod
    def parse_file(file_path: str, file_name: str) -> str:
        """
        根据扩展名选择解析器，返回纯文本

        Args:
            file_path: 本地文件路径
            file_name: 原始文件名

        Returns:
            提取的纯文本
        """
        ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""

        if ext == "pdf":
            return DocumentParser._parse_pdf(file_path)
        elif ext in ("doc", "docx"):
            return DocumentParser._parse_docx(file_path)
        elif ext in ("xls", "xlsx"):
            return DocumentParser._parse_xlsx(file_path)
        elif ext == "txt":
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _parse_pdf(path: str) -> str:
        """PyMuPDF 提取 PDF 文本，回退 pdfplumber"""
        import fitz

        texts = []
        doc = fitz.open(path)
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                texts.append(text.strip())
        doc.close()

        result = "\n".join(texts)

        # 如果 PyMuPDF 提取文本太少，回退 pdfplumber
        if len(result.strip()) < 100:
            try:
                import pdfplumber

                texts2 = []
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            texts2.append(text)
                fallback = "\n".join(texts2)
                if len(fallback.strip()) > len(result.strip()):
                    result = fallback
            except Exception:
                pass

        return result

    @staticmethod
    def pdf_to_images_base64(
        path: str, max_pages: int = 5, dpi: int = 150
    ) -> list[str]:
        """
        将 PDF 页面渲染为图片并返回 base64 编码列表。
        用于处理扫描件/图片型 PDF。
        """
        import fitz
        import base64

        images = []
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode()
            images.append(b64)
        doc.close()
        return images

    @staticmethod
    def is_image_pdf(path: str) -> bool:
        """判断 PDF 是否为扫描件/图片型（文本内容极少）"""
        import fitz

        doc = fitz.open(path)
        page_count = len(doc)
        if page_count == 0:
            doc.close()
            return False

        total_text = ""
        for page in doc:
            total_text += page.get_text()
        doc.close()

        cleaned = "".join(c for c in total_text if c.strip())
        # 总文字少于 200 字符，或每页平均少于 100 字符，视为图片 PDF
        # 网页截图型 PDF 虽然总字数多，但每页文字很少
        avg_per_page = len(cleaned) / page_count
        return len(cleaned) < 200 or avg_per_page < 100

    @staticmethod
    def _parse_docx(path: str) -> str:
        """python-docx 提取 Word 文本，含文本框/Shape"""
        from docx import Document

        doc = Document(path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    texts.append(row_text)

        # 普通段落/表格没提取到文本时，尝试从 XML 提取文本框（txbxContent）内容
        # 品牌设计类 DOCX 常把文字放在 Word 文本框/Shape 里，doc.paragraphs 读不到
        if not texts:
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for txbx in doc.element.body.findall(".//w:txbxContent", ns):
                for p in txbx.findall(".//w:p", ns):
                    text = "".join(t.text or "" for t in p.findall(".//w:t", ns))
                    if text.strip():
                        texts.append(text)
            if texts:
                logger.info(f"_parse_docx: 从文本框提取到 {len(texts)} 段文本")

        return "\n".join(texts)

    @staticmethod
    def _parse_xlsx(path: str) -> str:
        """openpyxl 提取 Excel 文本（所有 sheet 拼接）"""
        from openpyxl import load_workbook

        def _extract(read_only: bool) -> str:
            wb = load_workbook(path, read_only=read_only, data_only=True)
            texts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        texts.append(row_text)
            wb.close()
            return "\n".join(texts)

        result = _extract(read_only=True)

        # 某些营销脚本 Excel 在 openpyxl 的只读模式下会漏掉绝大多数行，
        # 表现为只能读到首行标题。检测到文本异常短时回退到普通模式重试。
        if len(result.strip()) < 50:
            fallback = _extract(read_only=False)
            if len(fallback.strip()) > len(result.strip()):
                logger.info(
                    "_parse_xlsx: fallback to normal mode, extracted %s -> %s chars",
                    len(result.strip()),
                    len(fallback.strip()),
                )
                result = fallback

        return result

    @staticmethod
    def _parse_txt(path: str) -> str:
        """纯文本文件"""
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def _extract_docx_images(path: str) -> list[str]:
        """从 DOCX 文件中提取嵌入图片（DOCX 本质是 ZIP，图片在 word/media/ 目录）"""
        import zipfile
        import base64

        images = []
        image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

        try:
            with zipfile.ZipFile(path, "r") as zf:
                for name in zf.namelist():
                    if not name.startswith("word/media/"):
                        continue
                    ext = os.path.splitext(name)[1].lower()
                    if ext not in image_exts:
                        continue
                    img_data = zf.read(name)
                    b64 = base64.b64encode(img_data).decode()
                    if len(b64) > DocumentParser.MAX_IMAGE_SIZE:
                        logger.debug(f"跳过过大图片: {name} ({len(b64)} bytes)")
                        continue
                    images.append(b64)
                    if len(images) >= DocumentParser.MAX_IMAGES:
                        break
        except Exception as e:
            logger.warning(f"提取 DOCX 图片失败: {e}")

        return images

    @staticmethod
    def _extract_xlsx_images(path: str) -> list[str]:
        """从 XLSX 文件中提取嵌入图片（通过 openpyxl 的 _images 属性）"""
        import base64

        images = []
        try:
            from openpyxl import load_workbook

            wb = load_workbook(path, read_only=False)
            for sheet in wb.worksheets:
                for img in getattr(sheet, "_images", []):
                    try:
                        img_data = img._data()
                        b64 = base64.b64encode(img_data).decode()
                        if len(b64) > DocumentParser.MAX_IMAGE_SIZE:
                            continue
                        images.append(b64)
                        if len(images) >= DocumentParser.MAX_IMAGES:
                            break
                    except Exception:
                        continue
                if len(images) >= DocumentParser.MAX_IMAGES:
                    break
            wb.close()
        except Exception as e:
            logger.warning(f"提取 XLSX 图片失败: {e}")

        return images
